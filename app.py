"""
app.py — GestiMuni: Sistema de Gestión Municipal con ML
Roles: admin | empleado | ciudadano
Incluye: subida de docs, exportación Excel, alertas por correo al ciudadano,
         búsqueda por DNI, CRUD completo de usuarios, dashboard ML
"""
import csv
import docx2txt
import random, os, io, uuid, json, re, zipfile, unicodedata
from flask import (Flask, render_template, redirect, url_for, request,
                   flash, session, jsonify, send_file, abort)
from flask_login import (LoginManager, login_user, logout_user,
                         login_required, current_user)
from flask_mail import Mail, Message
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from models import db, Usuario, Tramite, Postulante, Alerta, Documento, LogActividad
from ml import (clasificar_tramite, evaluar_cv,
                obtener_stats_tramites, obtener_stats_cvs, reentrenar, explicar_tramite)
from datetime import datetime
from functools import wraps
import random, os, io, uuid

# ── ExcelJS equivalente en Python: openpyxl ──────────────────────────
from openpyxl import Workbook
from openpyxl.styles import (Font, PatternFill, Alignment,
                              Border, Side, GradientFill)
from openpyxl.utils import get_column_letter
from services.auditoria_service import registrar_auditoria
from models import db, Usuario, Tramite, Postulante, Alerta, Auditoria
# ── dotenv ───────────────────────────────────────────────────────────
# ver documentos
from openpyxl import load_workbook
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from dotenv import load_dotenv
load_dotenv()

# ════════════════════════════════════════════════════════════════════
app = Flask(__name__)
app.config['SECRET_KEY']                  = os.getenv('SECRET_KEY', 'gestimuni_dev_2024')
app.config['SQLALCHEMY_DATABASE_URI']     = os.getenv('SQLALCHEMY_DATABASE_URI',
                                                       'sqlite:///gestimuni.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAIL_SERVER']                 = os.getenv('MAIL_SERVER', 'smtp.gmail.com')
app.config['MAIL_PORT']                   = int(os.getenv('MAIL_PORT', 587))
app.config['MAIL_USE_TLS']                = os.getenv('MAIL_USE_TLS', 'True') == 'True'
app.config['MAIL_USERNAME']               = os.getenv('MAIL_USERNAME', '')
app.config['MAIL_PASSWORD']               = os.getenv('MAIL_PASSWORD', '')
app.config['MAIL_DEFAULT_SENDER']         = os.getenv('MAIL_USERNAME', '')
app.config['UPLOAD_FOLDER']               = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['MAX_CONTENT_LENGTH']          = 16 * 1024 * 1024   # 16 MB

EXTENSIONES_PERMITIDAS = {'pdf', 'doc', 'docx', 'xls', 'xlsx',
                           'png', 'jpg', 'jpeg', 'gif', 'txt'}

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

db.init_app(app)
mail        = Mail(app)
login_mgr   = LoginManager(app)
login_mgr.login_view = 'login'

# ── Helpers ──────────────────────────────────────────────────────────
@login_mgr.user_loader
def load_user(uid):
    return Usuario.query.get(int(uid))


def extension_permitida(filename):
    return ('.' in filename and
            filename.rsplit('.', 1)[1].lower() in EXTENSIONES_PERMITIDAS)


def rol_requerido(*roles):
    def decorator(f):
        @wraps(f)
        def wrapped(*args, **kwargs):
            if current_user.rol not in roles:
                flash('No tienes permiso para acceder.', 'danger')
                return redirect(url_for('dashboard'))
            return f(*args, **kwargs)
        return wrapped
    return decorator


def registrar_log(accion, entidad=None, entidad_id=None):
    log = LogActividad(
        usuario_id=current_user.id,
        accion=accion,
        entidad=entidad,
        entidad_id=entidad_id,
        ip=request.remote_addr
    )
    db.session.add(log)


def enviar_correo_ciudadano(usuario, asunto, cuerpo_html):
    """Envía correo al email real del ciudadano si lo tiene registrado."""

    if not usuario.email:
        return False

    try:
        msg = Message(
            subject=f'GestiMuni — {asunto}',
            sender=f"GestiMuni <{app.config['MAIL_USERNAME']}>",
            recipients=[usuario.email],
            html=cuerpo_html
        )

        mail.send(msg)
        print(f"Correo enviado correctamente a {usuario.email}")
        return True

    except Exception as e:
        app.logger.error(f'Error enviando correo a {usuario.email}: {e}')
        print(f"Error enviando correo a {usuario.email}: {e}")
        return False

def _html_alerta_tramite(ciudadano_nombre, tramite_tipo, tramite_estado,
                          tramite_id, mensaje_extra=''):
    colores = {
        'aprobado':  ('#1D9E75', '#E1F5EE'),
        'rechazado': ('#A32D2D', '#FCEBEB'),
        'en_proceso':('#185FA5', '#E6F1FB'),
        'pendiente': ('#854F0B', '#FAEEDA'),
    }
    color_btn, color_bg = colores.get(tramite_estado, ('#444', '#f5f5f5'))
    return f'''
    <div style="font-family:Arial,sans-serif;max-width:520px;margin:0 auto;
                border:1px solid #dee2e6;border-radius:10px;overflow:hidden;">
      <div style="background:#2E7D4F;padding:18px 24px;text-align:center;">
        <h2 style="color:#fff;margin:0;font-size:20px;">🏛 GestiMuni Huánuco</h2>
        <p style="color:#a8d5ba;margin:4px 0 0;font-size:13px;">
          Municipalidad Provincial de Yau</p>
      </div>
      <div style="padding:24px;">
        <p style="color:#333;font-size:15px;">Hola <strong>{ciudadano_nombre}</strong>,</p>
        <p style="color:#555;font-size:14px;">
          Tu trámite <strong>#{tramite_id}</strong>
          (<em>{tramite_tipo.replace('_',' ').title()}</em>)
          ha sido actualizado.</p>
        <div style="background:{color_bg};border-radius:8px;padding:14px 18px;
                    margin:16px 0;text-align:center;">
          <span style="color:{color_btn};font-size:18px;font-weight:700;
                        letter-spacing:1px;">
            {tramite_estado.upper().replace('_',' ')}
          </span>
        </div>
        {f'<p style="color:#555;font-size:13px;">{mensaje_extra}</p>' if mensaje_extra else ''}
        <p style="color:#888;font-size:12px;margin-top:20px;">
          Puedes ingresar al sistema para más detalles.<br>
          Este correo es generado automáticamente, no respondas a él.</p>
      </div>
      <div style="background:#f8f9fa;padding:12px 24px;text-align:center;">
        <p style="color:#aaa;font-size:11px;margin:0;">
          © 2024 GestiMuni — Sistema de Gestión Municipal con IA</p>
      </div>
    </div>'''


# ════════════════════════════════════════════════════════════════════
# RUTAS PÚBLICAS
# ════════════════════════════════════════════════════════════════════
@app.route('/')
def index():
    return redirect(url_for('inicio'))


@app.route('/inicio')
def inicio():
    return render_template('inicio.html')


# ── Registro ciudadano ───────────────────────────────────────────────
@app.route('/registro', methods=['GET', 'POST'])
def registro():
    if request.method == 'POST':
        dni      = request.form.get('dni', '').strip()
        nombre   = request.form['nombre'].strip()
        username = request.form['username'].strip()
        email    = request.form.get('email', '').strip() or None
        telefono = request.form.get('telefono', '').strip() or None
        password = request.form['password']

        if Usuario.query.filter_by(username=username).first():
            flash('El usuario ya existe.', 'danger')
            return redirect(url_for('registro'))
        if dni and Usuario.query.filter_by(dni=dni).first():
            flash('El DNI ya está registrado.', 'danger')
            return redirect(url_for('registro'))

        usuario = Usuario(nombre=nombre, username=username,
                          password=generate_password_hash(password),
                          rol='ciudadano', dni=dni or None,
                          email=email, telefono=telefono)
        db.session.add(usuario)
        db.session.commit()
        flash('Cuenta creada. Inicia sesión.', 'success')
        return redirect(url_for('login'))
    return render_template('registro.html')


# ── Registro de postulante a puestos de trabajo ─────────────────────
@app.route('/registro/postulante', methods=['GET', 'POST'])
def registro_postulante():
    if request.method == 'POST':
        dni = request.form.get('dni', '').strip()
        nombre = request.form.get('nombre', '').strip()
        username = request.form.get('username', '').strip()
        email = request.form.get('email', '').strip() or None
        telefono = request.form.get('telefono', '').strip() or None
        password = request.form.get('password', '')

        if not re.fullmatch(r'\d{8}', dni):
            flash('El DNI debe contener exactamente 8 dígitos.', 'warning')
            return redirect(url_for('registro_postulante'))
        if Usuario.query.filter_by(username=username).first():
            flash('El usuario ya existe.', 'danger')
            return redirect(url_for('registro_postulante'))
        if Usuario.query.filter_by(dni=dni).first():
            flash('El DNI ya está registrado.', 'danger')
            return redirect(url_for('registro_postulante'))

        usuario = Usuario(
            nombre=nombre, username=username,
            password=generate_password_hash(password),
            rol='postulante', dni=dni, email=email, telefono=telefono
        )
        db.session.add(usuario)
        db.session.commit()
        flash('Cuenta de postulante creada. Inicia sesión para subir tu CV.', 'success')
        return redirect(url_for('login'))

    return render_template('registro_postulante.html')


# ── Login ciudadano ──────────────────────────────────────────────────
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user = Usuario.query.filter_by(username=username).first()
        if user and check_password_hash(user.password, password) and user.activo:
            if user.rol in ('ciudadano', 'empleado', 'postulante'):
                login_user(user)
                user.ultimo_acceso = datetime.utcnow()
                db.session.commit()
                registrar_auditoria(
                    f'Inicio de sesión de {user.nombre}',
                    'Seguridad'
                )
                if user.rol == 'postulante':
                    return redirect(url_for('postulante_panel'))
                
                return redirect(url_for('dashboard'))
            else:
                flash('Usa el acceso de administrador.', 'danger')
        else:
            flash('Usuario o contraseña incorrectos.', 'danger')
    return render_template('login.html')


# ── Login admin con PIN ──────────────────────────────────────────────
@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        paso = request.form.get('paso')
        if paso == '1':
            username = request.form['username']
            password = request.form['password']
            user = Usuario.query.filter_by(username=username).first()
            if user and check_password_hash(user.password, password):
                if user.rol == 'admin':
                    pin = str(random.randint(100000, 999999))
                    session['admin_pin']      = pin
                    session['admin_username'] = username
                    try:
                        msg = Message(
                            subject='GestiMuni — Código de verificación',
                            recipients=[user.email or app.config['MAIL_USERNAME']],
                            html=f'''
                            <div style="font-family:Arial;max-width:400px;margin:auto;
                                        border:1px solid #dee2e6;border-radius:8px;">
                              <div style="background:#2E7D4F;padding:16px;
                                          text-align:center;border-radius:8px 8px 0 0;">
                                <h2 style="color:#fff;margin:0;">GestiMuni Huánuco</h2>
                              </div>
                              <div style="padding:24px;text-align:center;">
                                <p style="color:#555;">Código de verificación admin:</p>
                                <div style="font-size:2.5rem;font-weight:700;
                                            letter-spacing:10px;color:#1B5E35;
                                            padding:15px;background:#E8F5EE;
                                            border-radius:6px;">{pin}</div>
                                <p style="color:#888;font-size:12px;margin-top:12px;">
                                  Expira en 5 minutos.</p>
                              </div>
                            </div>''')
                        mail.send(msg)
                        flash('Código enviado a tu correo.', 'success')
                    except Exception as e:
                        flash(f'Error al enviar correo: {e}', 'danger')
                        return render_template('admin_login.html', paso=1)
                    return render_template('admin_login.html', paso=2)
                else:
                    flash('No tienes permiso de administrador.', 'danger')
            else:
                flash('Usuario o contraseña incorrectos.', 'danger')

        elif paso == '2':
            pin_ok   = session.get('admin_pin')
            username = session.get('admin_username')
            if request.form['pin'] == pin_ok:
                user = Usuario.query.filter_by(username=username).first()
                login_user(user)
                user.ultimo_acceso = datetime.utcnow()
                db.session.commit()
                session.pop('admin_pin', None)
                session.pop('admin_username', None)
                registrar_log('login_admin')
                db.session.commit()
                registrar_auditoria(
                    f'Inicio de sesión administrador de {user.nombre}',
                    'Seguridad'
                )
                return redirect(url_for('dashboard'))
            else:
                flash('Código incorrecto.', 'danger')
                return render_template('admin_login.html', paso=2)

    return render_template('admin_login.html', paso=1)


# ── Registro de nuevo admin ──────────────────────────────────────────
@app.route('/admin/registro', methods=['GET', 'POST'])
@login_required
@rol_requerido('admin')
def admin_registro():
    if request.method == 'POST':
        nombre   = request.form['nombre'].strip()
        username = request.form['username'].strip()
        email    = request.form.get('email', '').strip() or None
        password = request.form['password']

        if Usuario.query.filter_by(username=username).first():
            flash('El usuario ya existe.', 'danger')
            return redirect(url_for('admin_registro'))

        usuario = Usuario(
            nombre=nombre,
            username=username,
            password=generate_password_hash(password),
            rol='admin',
            email=email
        )
        db.session.add(usuario)
        registrar_auditoria(f'Creó nuevo admin: {nombre}', 'Usuarios')
        db.session.commit()
        flash(f'Admin "{username}" creado correctamente.', 'success')
        return redirect(url_for('dashboard'))
    return render_template('admin_registro.html')


@app.route('/logout')
@login_required
def logout():
    registrar_auditoria(
        f'Cierre de sesión de {current_user.nombre}',
        'Seguridad'
    )
    registrar_log('logout')
    db.session.commit()
    logout_user()
    return redirect(url_for('inicio'))


# ════════════════════════════════════════════════════════════════════
# DASHBOARD
# ════════════════════════════════════════════════════════════════════
@app.route('/dashboard')
@login_required
def dashboard():
    procesar_convocatorias_vencidas()
    if current_user.rol == 'admin':
        total   = Tramite.query.count()
        criticos= Tramite.query.filter_by(prioridad_ml='critico').count()
        postulantes = Postulante.query.count()
        alertas_n= Alerta.query.filter_by(leido=False).count()
        tramites = Tramite.query.order_by(Tramite.fecha_registro.desc()).limit(5).all()
        return render_template('dashboard.html',
            total_tramites=total, criticos=criticos,
            postulantes=postulantes, alertas=alertas_n, tramites=tramites)

    elif current_user.rol == 'empleado':
        mis   = Tramite.query.filter_by(empleado_id=current_user.id).all()
        pend  = Tramite.query.filter_by(estado='pendiente').count()
        return render_template('dashboard_empleado.html',
            tramites=mis, pendientes=pend)

    elif current_user.rol == 'postulante':
        return redirect(url_for('postulante_panel'))

    else:  # ciudadano
        mis_tramites = Tramite.query.filter_by(ciudadano_id=current_user.id).all()
        mis_alertas  = Alerta.query.filter_by(
            ciudadano_id=current_user.id, leido=False).all()
        return render_template('ciudadano.html',
            tramites=mis_tramites, alertas=mis_alertas)


# ════════════════════════════════════════════════════════════════════
# TRÁMITES
# ════════════════════════════════════════════════════════════════════
@app.route('/tramites')
@login_required
def tramites():
    if current_user.rol == 'admin':
        lista = Tramite.query.order_by(Tramite.fecha_registro.desc()).all()
    elif current_user.rol == 'empleado':
        lista = Tramite.query.order_by(Tramite.fecha_registro.desc()).all()
    else:
        lista = Tramite.query.filter_by(
            ciudadano_id=current_user.id).order_by(Tramite.fecha_registro.desc()).all()
    return render_template('tramites.html', tramites=lista)


@app.route('/tramites/nuevo', methods=['GET', 'POST'])
@login_required
def nuevo_tramite():
    if request.method == 'POST':
        tipo        = request.form['tipo']
        descripcion = request.form['descripcion']
        urgencias_map = {
            'denuncia_peligro':        'alta',
            'licencia_construccion':   'media',
            'licencia_funcionamiento': 'media',
            'permiso_negocio':         'media',
            'certificado_residencia':  'baja',
            'partida_nacimiento':      'baja'
        }
        urgencia  = urgencias_map.get(tipo, 'media')
        
        dias_resolucion = int(request.form.get('dias_resolucion', 5))
        cantidad_documentos = int(request.form.get('cantidad_documentos', 1))
        
        prioridad, confianza = clasificar_tramite(tipo, urgencia, dias_resolucion, cantidad_documentos)
        explicacion = explicar_tramite(tipo, urgencia, prioridad, dias_resolucion, cantidad_documentos)

        tramite = Tramite(
            ciudadano_id=current_user.id,
            tipo=tipo, 
            descripcion=descripcion,
            urgencia=urgencia, 
            prioridad_ml=prioridad,
            confianza_ml=confianza, 
            dias_resolucion=dias_resolucion,
            explicacion_ml=explicacion
            
        )
        db.session.add(tramite)
        db.session.flush()

        # Alerta interna
        alerta = Alerta(
            tramite_id=tramite.id,
            ciudadano_id=current_user.id,
            mensaje=(f'Tu trámite "{tipo.replace("_"," ").title()}" fue registrado '
                     f'con prioridad {prioridad.upper()} '
                     f'(confianza ML: {round(confianza*100,1)}%).'),
            tipo='info'
        )
        db.session.add(alerta)

        # Subida de documentos adjuntos
        archivos = request.files.getlist('documentos')
        for archivo in archivos:
            if archivo and extension_permitida(archivo.filename):
                nombre_orig = secure_filename(archivo.filename)
                ext         = nombre_orig.rsplit('.', 1)[1].lower()
                nombre_disco= f"{uuid.uuid4().hex}.{ext}"
                ruta        = os.path.join(app.config['UPLOAD_FOLDER'], nombre_disco)
                archivo.save(ruta)
                tam = os.path.getsize(ruta)
                doc = Documento(
                    tramite_id=tramite.id,
                    usuario_id=current_user.id,
                    nombre_original=nombre_orig,
                    nombre_archivo=nombre_disco,
                    tipo_mime=archivo.content_type,
                    tamanio=tam
                )
                db.session.add(doc)

        registrar_log('nuevo_tramite', 'Tramite', tramite.id)
        db.session.commit()
        
        registrar_auditoria(
            f"Registró trámite #{tramite.id}: {tipo.replace('_', ' ').title()} con prioridad {prioridad.upper()}",
            "Trámites"
        )
        

        # Correo al ciudadano
        html_c = _html_alerta_tramite(
            current_user.nombre, tipo, 'pendiente', tramite.id,
            f'Prioridad asignada por IA: {prioridad.upper()}'
        )
        enviado = enviar_correo_ciudadano(
            current_user, 'Trámite registrado', html_c)
        if enviado:
            alerta.enviado_email = True
            db.session.commit()

        flash(f'Trámite registrado. Prioridad IA: {prioridad.upper()} '
              f'(confianza {round(confianza*100,1)}%)', 'success')
        return redirect(url_for('tramites'))
    return render_template('nuevo_tramite.html')


@app.route('/tramites/estado/<int:id>', methods=['POST'])
@login_required
@rol_requerido('admin', 'empleado')
def cambiar_estado(id):
    tramite = Tramite.query.get_or_404(id)
    nuevo_estado  = request.form['estado']
    observaciones = request.form.get('observaciones', '')
    tramite.estado        = nuevo_estado
    tramite.observaciones = observaciones
    tramite.fecha_actualizacion = datetime.utcnow()

    alerta = Alerta(
        tramite_id=tramite.id,
        ciudadano_id=tramite.ciudadano_id,
        mensaje=(f'Tu trámite "{tramite.tipo.replace("_"," ").title()}" '
                 f'cambió a: {nuevo_estado.upper().replace("_"," ")}.'
                 + (f' Nota: {observaciones}' if observaciones else '')),
        tipo='success' if nuevo_estado == 'aprobado' else
             'danger'  if nuevo_estado == 'rechazado' else 'info'
    )
    db.session.add(alerta)
    registrar_log('cambiar_estado', 'Tramite', id)
    db.session.commit()
    registrar_auditoria(
        f"Cambió el trámite #{tramite.id} al estado {nuevo_estado.upper().replace('_', ' ')}",
        "Trámites"
    )

    # Correo al ciudadano
    html_c = _html_alerta_tramite(
        tramite.ciudadano.nombre, tramite.tipo,
        nuevo_estado, tramite.id, observaciones
    )
    enviado = enviar_correo_ciudadano(tramite.ciudadano,
                                      f'Estado actualizado: {nuevo_estado}', html_c)
    if enviado:
        alerta.enviado_email = True
        db.session.commit()

    flash('Estado actualizado.', 'success')
    return redirect(url_for('tramites'))


@app.route('/tramites/<int:id>')
@login_required
def detalle_tramite(id):
    tramite = Tramite.query.get_or_404(id)
    if (current_user.rol == 'ciudadano' and
            tramite.ciudadano_id != current_user.id):
        abort(403)
    return render_template('detalle_tramite.html', tramite=tramite)


# ════════════════════════════════════════════════════════════════════
# DOCUMENTOS
# ════════════════════════════════════════════════════════════════════
@app.route('/documentos/ver/<int:doc_id>')
@login_required
def ver_documento(doc_id):
    doc = Documento.query.get_or_404(doc_id)
    ruta = os.path.join(app.config['UPLOAD_FOLDER'], doc.nombre_archivo)

    if not os.path.exists(ruta):
        abort(404)

    extension = doc.nombre_original.rsplit('.', 1)[-1].lower()

    contenido_preview = None
    tipo_preview = extension

    if extension in ['pdf', 'png', 'jpg', 'jpeg', 'gif', 'txt']:
        return send_file(
            ruta,
            mimetype=doc.tipo_mime,
            as_attachment=False,
            download_name=doc.nombre_original
        )

    if extension in ['xlsx', 'xls']:
        wb = load_workbook(ruta, data_only=True)
        ws = wb.active

        filas = []
        for row in ws.iter_rows(values_only=True):
            filas.append([cell if cell is not None else '' for cell in row])

        contenido_preview = filas

    elif extension in ['docx']:
        documento = DocxDocument(ruta)
        contenido_preview = [p.text for p in documento.paragraphs if p.text.strip()]

    else:
        flash('Vista previa no disponible para este tipo de archivo.', 'warning')
        return redirect(url_for('detalle_tramite', id=doc.tramite_id))

    return render_template(
        'preview_documento.html',
        doc=doc,
        contenido=contenido_preview,
        tipo=tipo_preview
    )

@app.route('/documentos/descargar/<int:doc_id>')
@login_required
def descargar_documento(doc_id):
    doc  = Documento.query.get_or_404(doc_id)
    ruta = os.path.join(app.config['UPLOAD_FOLDER'], doc.nombre_archivo)
    if not os.path.exists(ruta):
        abort(404)
    return send_file(ruta, as_attachment=True,
                     download_name=doc.nombre_original)


@app.route('/documentos/eliminar/<int:doc_id>', methods=['POST'])
@login_required
@rol_requerido('admin', 'empleado')
def eliminar_documento(doc_id):
    doc  = Documento.query.get_or_404(doc_id)
    ruta = os.path.join(app.config['UPLOAD_FOLDER'], doc.nombre_archivo)
    if os.path.exists(ruta):
        os.remove(ruta)
    db.session.delete(doc)
    registrar_log('eliminar_doc', 'Documento', doc_id)
    db.session.commit()
    registrar_auditoria(
        f"Eliminó documento {doc.nombre_original}",
        "Documentos"
    )
    flash('Documento eliminado.', 'success')
    return redirect(request.referrer or url_for('tramites'))


# ════════════════════════════════════════════════════════════════════
# CURRÍCULOS / POSTULANTES
# ════════════════════════════════════════════════════════════════════
@app.route('/curriculos')
@login_required
@rol_requerido('admin', 'empleado')
def curriculos():
    procesar_convocatorias_vencidas()
    lista = Postulante.query.order_by(Postulante.puesto.asc(), Postulante.puntaje_ml.desc()).all()
    return render_template('curriculos.html', postulantes=lista, puestos=cargar_puestos())


def convocatorias_abiertas():
    ahora = datetime.now()
    abiertas = []
    for conv in cargar_convocatorias():
        if conv.get('finalizada'):
            continue
        inicio = _parse_dt(conv.get('fecha_inicio'))
        cierre = _parse_dt(conv.get('fecha_cierre'))
        if (not inicio or ahora >= inicio) and (not cierre or ahora <= cierre):
            abiertas.append(conv)
    return abiertas


@app.route('/postulante/cv', methods=['GET', 'POST'])
@login_required
@rol_requerido('postulante')
def postulante_panel():
    puestos = cargar_puestos()
    abiertas = convocatorias_abiertas()
    postulaciones = Postulante.query.filter_by(dni=current_user.dni).order_by(Postulante.fecha_registro.desc()).all()

    if request.method == 'POST':
        puesto_key = request.form.get('puesto', '').strip()
        email_form = request.form.get('email', '').strip() or current_user.email
        archivo_cv = request.files.get('archivo_cv')

        puestos_abiertos = {c.get('puesto') for c in abiertas}
        if puesto_key not in puestos or puesto_key not in puestos_abiertos:
            flash('Seleccione un puesto con convocatoria abierta.', 'warning')
            return redirect(url_for('postulante_panel'))
        if not archivo_cv or not archivo_cv.filename:
            flash('Debe subir su currículum vitae.', 'warning')
            return redirect(url_for('postulante_panel'))

        nombre_original = secure_filename(archivo_cv.filename)
        extension = nombre_original.rsplit('.', 1)[-1].lower()
        if extension not in ['pdf', 'docx', 'txt', 'csv']:
            flash('Formato no permitido. Use PDF, DOCX, TXT o CSV.', 'danger')
            return redirect(url_for('postulante_panel'))

        nombre_archivo_cv = f"cv_{uuid.uuid4().hex}.{extension}"
        ruta_cv = os.path.join(app.config['UPLOAD_FOLDER'], nombre_archivo_cv)
        archivo_cv.save(ruta_cv)

        texto_cv = extraer_texto_cv(ruta_cv, extension)
        if not texto_cv:
            flash('No se pudo leer el CV. Verifique que el archivo tenga texto legible.', 'danger')
            return redirect(url_for('postulante_panel'))

        analisis = analizar_cv_por_puesto(texto_cv, puesto_key)

        postulante = Postulante(
            nombre=current_user.nombre,
            dni=current_user.dni,
            email=email_form,
            puesto=puesto_key,
            experiencia=analisis['experiencia'],
            educacion=analisis['educacion'],
            habilidades=analisis['habilidades'],
            idiomas=str(analisis['certificaciones']),
            resultado_ml=analisis['resultado'],
            puntaje_ml=analisis['puntaje'],
            probabilidad=analisis['probabilidad'],
            archivo_cv=nombre_archivo_cv,
            texto_cv=analisis['explicacion']
        )
        db.session.add(postulante)
        if email_form and email_form != current_user.email:
            current_user.email = email_form
        registrar_log('postulacion_cv', 'Postulante')
        registrar_auditoria(
            f"Postulante {current_user.nombre} subió CV para {puestos[puesto_key]['nombre']}",
            'Postulaciones'
        )
        db.session.commit()

        flash('Gracias por postular. Recibirás un correo cuando finalice la convocatoria.', 'success')
        return redirect(url_for('postulante_panel'))

    return render_template(
        'postulante_panel.html',
        puestos=puestos,
        convocatorias=abiertas,
        postulaciones=postulaciones
    )

@app.route('/curriculos/nuevo', methods=['GET', 'POST'])
@login_required
@rol_requerido('admin', 'empleado')
def nuevo_curriculo():
    puestos = cargar_puestos()

    if request.method == 'POST':
        nombre_form = request.form.get('nombre', '').strip()
        dni_form = request.form.get('dni', '').strip() or None
        email_form = request.form.get('email', '').strip() or None
        puesto_key = request.form.get('puesto', '').strip()
        archivo_cv = request.files.get('archivo_cv')

        if dni_form and not re.fullmatch(r'\d{8}', dni_form):
            flash('El DNI debe contener exactamente 8 dígitos.', 'warning')
            return redirect(url_for('nuevo_curriculo'))

        if not puesto_key or puesto_key not in puestos:
            flash('Seleccione un puesto válido.', 'warning')
            return redirect(url_for('nuevo_curriculo'))

        if not archivo_cv or not archivo_cv.filename:
            flash('Debe subir el currículum vitae en PDF, DOCX, TXT o CSV.', 'warning')
            return redirect(url_for('nuevo_curriculo'))

        nombre_original = secure_filename(archivo_cv.filename)
        extension = nombre_original.rsplit('.', 1)[-1].lower()

        if extension not in ['pdf', 'docx', 'txt', 'csv']:
            flash('Formato no permitido. Use PDF, DOCX, TXT o CSV.', 'danger')
            return redirect(url_for('nuevo_curriculo'))

        nombre_archivo_cv = f"cv_{uuid.uuid4().hex}.{extension}"
        ruta_cv = os.path.join(app.config['UPLOAD_FOLDER'], nombre_archivo_cv)
        archivo_cv.save(ruta_cv)

        texto_cv = extraer_texto_cv(ruta_cv, extension)

        if not texto_cv:
            flash('No se pudo extraer texto del CV. Verifique que el archivo tenga contenido legible.', 'danger')
            return redirect(url_for('nuevo_curriculo'))

        analisis = analizar_cv_por_puesto(texto_cv, puesto_key)

        nombre_final = nombre_form or analisis.get('nombre_detectado') or 'Postulante no identificado'
        dni_final = dni_form or analisis.get('dni_detectado')
        email_final = email_form or analisis.get('correo_detectado')

        postulante = Postulante(
            nombre=nombre_final,
            dni=dni_final,
            email=email_final,
            puesto=puesto_key,
            experiencia=analisis['experiencia'],
            educacion=analisis['educacion'],
            habilidades=analisis['habilidades'],
            idiomas=str(analisis['certificaciones']),
            resultado_ml=analisis['resultado'],
            puntaje_ml=analisis['puntaje'],
            probabilidad=analisis['probabilidad'],
            archivo_cv=nombre_archivo_cv,
            texto_cv=analisis['explicacion']
        )

        db.session.add(postulante)
        registrar_log('nuevo_curriculo', 'Postulante')
        registrar_auditoria(
            f"Evaluó CV de {nombre_final} para el puesto {puestos[puesto_key]['nombre']}: "
            f"{analisis['resultado'].upper()} con puntaje {analisis['puntaje']}%",
            'Currículos'
        )
        db.session.commit()

        flash(
            f"CV evaluado: {analisis['resultado'].upper()} — Puntaje: {analisis['puntaje']}% — "
            f"Recomendación: {analisis['recomendacion']}",
            'success'
        )
        return redirect(url_for('curriculos'))

    return render_template('nuevo_curriculo.html', puestos=puestos)


@app.route('/curriculos/eliminar/<int:id>', methods=['POST'])
@login_required
@rol_requerido('admin')
def eliminar_postulante(id):
    p = Postulante.query.get_or_404(id)
    db.session.delete(p)
    registrar_log('eliminar_postulante', 'Postulante', id)
    db.session.commit()
    registrar_auditoria(
        f"Eliminó postulante {p.nombre}",
        "Currículos"
    )
    flash('Postulante eliminado.', 'success')
    return redirect(url_for('curriculos'))


# ════════════════════════════════════════════════════════════════════
# USUARIOS — CRUD COMPLETO
# ════════════════════════════════════════════════════════════════════
@app.route('/usuarios')
@login_required
@rol_requerido('admin')
def usuarios():
    q    = request.args.get('q', '').strip()
    rol  = request.args.get('rol', '')
    query= Usuario.query
    if q:
        query = query.filter(
            db.or_(Usuario.nombre.ilike(f'%{q}%'),
                   Usuario.username.ilike(f'%{q}%'),
                   Usuario.dni.ilike(f'%{q}%')))
    if rol:
        query = query.filter_by(rol=rol)
    lista = query.order_by(Usuario.fecha_registro.desc()).all()
    return render_template('usuarios.html', usuarios=lista,
                           q=q, rol_filtro=rol)


@app.route('/usuarios/buscar-dni')
@login_required
@rol_requerido('admin', 'empleado')
def buscar_por_dni():
    dni = request.args.get('dni', '').strip()
    if not dni:
        return jsonify({'error': 'DNI vacío'}), 400
    user = Usuario.query.filter_by(dni=dni).first()
    if not user:
        return jsonify({'error': 'No encontrado'}), 404
    tramites = Tramite.query.filter_by(ciudadano_id=user.id).all()
    return jsonify({
        'id':       user.id,
        'nombre':   user.nombre,
        'dni':      user.dni,
        'email':    user.email,
        'telefono': user.telefono,
        'rol':      user.rol,
        'tramites': [{'id': t.id, 'tipo': t.tipo,
                      'estado': t.estado, 'prioridad': t.prioridad_ml}
                     for t in tramites]
    })


@app.route('/usuarios/nuevo', methods=['GET', 'POST'])
@login_required
@rol_requerido('admin')
def nuevo_usuario():
    if request.method == 'POST':
        dni      = request.form.get('dni', '').strip() or None
        nombre   = request.form['nombre'].strip()
        username = request.form['username'].strip()
        email    = request.form.get('email', '').strip() or None
        telefono = request.form.get('telefono', '').strip() or None
        password = request.form['password']
        rol      = request.form.get('rol', 'ciudadano')

        if Usuario.query.filter_by(username=username).first():
            flash('El usuario ya existe.', 'danger')
            return redirect(url_for('nuevo_usuario'))
        if dni and Usuario.query.filter_by(dni=dni).first():
            flash('El DNI ya está registrado.', 'danger')
            return redirect(url_for('nuevo_usuario'))

        usuario = Usuario(nombre=nombre, username=username,
                          password=generate_password_hash(password),
                          rol=rol, dni=dni, email=email, telefono=telefono)
        db.session.add(usuario)
        registrar_log('crear_usuario', 'Usuario')
        db.session.commit()
        registrar_auditoria(
            f'Creó usuario {nombre} con rol {rol}',
            'Usuarios'
        )
        flash('Usuario creado correctamente.', 'success')
        return redirect(url_for('usuarios'))
    return render_template('nuevo_usuario.html')


@app.route('/usuarios/editar/<int:id>', methods=['GET', 'POST'])
@login_required
@rol_requerido('admin')
def editar_usuario(id):
    usuario = Usuario.query.get_or_404(id)
    if request.method == 'POST':
        usuario.nombre   = request.form['nombre'].strip()
        usuario.email    = request.form.get('email', '').strip() or None
        usuario.telefono = request.form.get('telefono', '').strip() or None
        usuario.rol      = request.form.get('rol', usuario.rol)
        usuario.dni      = request.form.get('dni', '').strip() or None
        nuevo_pass       = request.form.get('password', '').strip()
        if nuevo_pass:
            usuario.password = generate_password_hash(nuevo_pass)
        registrar_log('editar_usuario', 'Usuario', id)
        db.session.commit()
        registrar_auditoria(
            f'Editó usuario {usuario.nombre}',
            'Usuarios'
        )
        flash('Usuario actualizado.', 'success')
        return redirect(url_for('usuarios'))
    return render_template('editar_usuario.html', usuario=usuario)


@app.route('/usuarios/toggle/<int:id>', methods=['POST'])
@login_required
@rol_requerido('admin')
def toggle_usuario(id):
    usuario = Usuario.query.get_or_404(id)
    if usuario.id == current_user.id:
        flash('No puedes desactivar tu propia cuenta.', 'danger')
        return redirect(url_for('usuarios'))
    usuario.activo = not usuario.activo
    accion = 'activar_usuario' if usuario.activo else 'desactivar_usuario'
    registrar_log(accion, 'Usuario', id)
    db.session.commit()
    registrar_auditoria(
        f'{"Activó" if usuario.activo else "Desactivó"} usuario {usuario.nombre}',
        'Usuarios'
    )
    flash(f'Usuario {"activado" if usuario.activo else "desactivado"}.', 'success')
    return redirect(url_for('usuarios'))


@app.route('/usuarios/eliminar/<int:id>', methods=['POST'])
@login_required
@rol_requerido('admin')
def eliminar_usuario(id):
    usuario = Usuario.query.get_or_404(id)
    if usuario.id == current_user.id:
        flash('No puedes eliminar tu propia cuenta.', 'danger')
        return redirect(url_for('usuarios'))
    # Reasignar trámites a anónimo antes de eliminar
    Tramite.query.filter_by(ciudadano_id=id).update({'ciudadano_id': current_user.id})
    db.session.delete(usuario)
    registrar_log('eliminar_usuario', 'Usuario', id)
    db.session.commit()
    registrar_auditoria(
        f"Eliminó usuario {usuario.nombre}",
        "Usuarios"
    )
    flash('Usuario eliminado.', 'success')
    return redirect(url_for('usuarios'))


# ════════════════════════════════════════════════════════════════════
# ALERTAS
# ════════════════════════════════════════════════════════════════════
@app.route('/alertas')
@login_required
def alertas():
    if current_user.rol == 'admin':
        lista = Alerta.query.order_by(Alerta.fecha.desc()).all()
    else:
        lista = Alerta.query.filter_by(
            ciudadano_id=current_user.id).order_by(Alerta.fecha.desc()).all()
    return render_template('alertas.html', alertas=lista)


@app.route('/alertas/marcar-leido/<int:id>', methods=['POST'])
@login_required
def marcar_leido(id):
    alerta = Alerta.query.get_or_404(id)
    alerta.leido = True
    db.session.commit()
    return jsonify({'ok': True})


@app.route('/alertas/marcar-todas', methods=['POST'])
@login_required
def marcar_todas_leidas():
    q = Alerta.query.filter_by(leido=False)
    if current_user.rol != 'admin':
        q = q.filter_by(ciudadano_id=current_user.id)
    q.update({'leido': True})
    db.session.commit()
    flash('Todas las alertas marcadas como leídas.', 'success')
    return redirect(url_for('alertas'))


# ════════════════════════════════════════════════════════════════════
# DASHBOARD DE ML — MÉTRICAS
# ════════════════════════════════════════════════════════════════════
@app.route('/ml/stats')
@login_required
@rol_requerido('admin')
def ml_stats():
    st_tramites = obtener_stats_tramites()
    st_cvs      = obtener_stats_cvs()
    return render_template('ml_stats.html',
                           st_tramites=st_tramites,
                           st_cvs=st_cvs)


@app.route('/ml/stats/json')
@login_required
@rol_requerido('admin')
def ml_stats_json():
    return jsonify({
        'tramites': obtener_stats_tramites(),
        'cvs':      obtener_stats_cvs()
    })


@app.route('/ml/reentrenar', methods=['POST'])
@login_required
@rol_requerido('admin')
def ml_reentrenar():
    resultado = reentrenar()
    registrar_log('reentrenar_modelos', 'ML')
    db.session.commit()
    registrar_auditoria(
        "Reentrenó los modelos de Machine Learning",
        "Machine Learning"
    )
    flash(f'Modelos reentrenados. Tramites accuracy='
          f'{resultado["tramites"].get("accuracy","?")} | '
          f'CVs accuracy={resultado["cvs"].get("accuracy","?")}', 'success')
    return redirect(url_for('ml_stats'))


# ════════════════════════════════════════════════════════════════════
# EXPORTAR EXCEL — openpyxl (equivalente a ExcelJS en Python)
# ════════════════════════════════════════════════════════════════════
VERDE     = 'FF2E7D4F'
VERDE_CLR = 'FFE1F5EE'
GRIS_HD   = 'FF444441'
GRIS_FIL  = 'FFF1EFE8'
BLANCO    = 'FFFFFFFF'

def _borde(grosor='thin'):
    s = Side(style=grosor)
    return Border(left=s, right=s, top=s, bottom=s)


def _encabezado_institucional(ws, titulo, subtitulo='Municipalidad Provincial de Yau'):
    """Crea cabecera institucional verde en la hoja."""
    ws.merge_cells('A1:H1')
    c = ws['A1']
    c.value          = 'GESTIMUNI — SISTEMA DE GESTIÓN MUNICIPAL CON IA'
    c.font            = Font(name='Calibri', bold=True, size=14, color=BLANCO)
    c.fill            = PatternFill('solid', fgColor=VERDE)
    c.alignment       = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 28

    ws.merge_cells('A2:H2')
    c = ws['A2']
    c.value           = subtitulo
    c.font            = Font(name='Calibri', size=10, color=BLANCO, italic=True)
    c.fill            = PatternFill('solid', fgColor='FF1B5E35')
    c.alignment       = Alignment(horizontal='center', vertical='center')

    ws.merge_cells('A3:H3')
    c = ws['A3']
    c.value           = titulo
    c.font            = Font(name='Calibri', bold=True, size=12, color=GRIS_HD)
    c.fill            = PatternFill('solid', fgColor=VERDE_CLR)
    c.alignment       = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[3].height = 22

    ws.merge_cells('A4:D4')
    ws['A4'].value = f'Fecha de generación: {datetime.now().strftime("%d/%m/%Y %H:%M")}'
    ws['A4'].font  = Font(name='Calibri', size=9, italic=True, color='FF888780')
    ws.row_dimensions[4].height = 16


def _fila_encabezado_tabla(ws, fila, columnas):
    for col_idx, titulo in enumerate(columnas, 1):
        c = ws.cell(row=fila, column=col_idx, value=titulo)
        c.font      = Font(name='Calibri', bold=True, size=10, color=BLANCO)
        c.fill      = PatternFill('solid', fgColor=VERDE)
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        c.border    = _borde()
    ws.row_dimensions[fila].height = 20


def _fila_datos(ws, fila, datos, alternado=False):
    fill_color = GRIS_FIL if alternado else BLANCO
    for col_idx, valor in enumerate(datos, 1):
        c = ws.cell(row=fila, column=col_idx, value=valor)
        c.font      = Font(name='Calibri', size=9)
        c.fill      = PatternFill('solid', fgColor=fill_color)
        c.alignment = Alignment(vertical='center', wrap_text=True)
        c.border    = _borde('thin')
    ws.row_dimensions[fila].height = 16


@app.route('/exportar/tramites')
@login_required
@rol_requerido('admin', 'empleado')
def exportar_tramites_excel():
    estado   = request.args.get('estado', '')
    prioridad= request.args.get('prioridad', '')

    query = Tramite.query
    if estado:
        query = query.filter_by(estado=estado)
    if prioridad:
        query = query.filter_by(prioridad_ml=prioridad)
    if current_user.rol == 'empleado':
        query = query.filter_by(empleado_id=current_user.id)
    lista = query.order_by(Tramite.fecha_registro.desc()).all()

    wb = Workbook()
    ws = wb.active
    ws.title = 'Trámites'

    _encabezado_institucional(ws,
        f'REPORTE DE TRÁMITES'
        + (f' — Estado: {estado.upper()}' if estado else '')
        + (f' — Prioridad: {prioridad.upper()}' if prioridad else ''))

    cols = ['N°', 'Ciudadano', 'DNI', 'Tipo de Trámite',
            'Urgencia', 'Prioridad ML', 'Estado', 'Confianza ML', 'Fecha']
    anchos = [5, 22, 10, 22, 10, 12, 12, 12, 16]
    for i, w in enumerate(anchos, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    _fila_encabezado_tabla(ws, 6, cols)

    COLORES_PRIORIDAD = {'critico': 'FFFCEBEB', 'normal': 'FFFAEEDA', 'bajo': 'FFE1F5EE'}
    for idx, t in enumerate(lista):
        fila = 7 + idx
        dni_c = t.ciudadano.dni or '—'
        datos = [
            idx + 1,
            t.ciudadano.nombre,
            dni_c,
            t.tipo.replace('_', ' ').title(),
            t.urgencia.upper(),
            t.prioridad_ml.upper(),
            t.estado.upper().replace('_', ' '),
            f'{round(t.confianza_ml * 100, 1)}%',
            t.fecha_registro.strftime('%d/%m/%Y %H:%M')
        ]
        _fila_datos(ws, fila, datos, idx % 2 == 1)
        # Color según prioridad en columna F
        color_p = COLORES_PRIORIDAD.get(t.prioridad_ml, BLANCO)
        ws.cell(fila, 6).fill = PatternFill('solid', fgColor=color_p)

    # Fila resumen
    fila_res = 7 + len(lista) + 1
    ws.merge_cells(f'A{fila_res}:G{fila_res}')
    ws[f'A{fila_res}'].value = f'Total: {len(lista)} trámites'
    ws[f'A{fila_res}'].font  = Font(bold=True, size=10)
    ws[f'A{fila_res}'].fill  = PatternFill('solid', fgColor=VERDE_CLR)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    nombre_archivo = f'tramites_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name=nombre_archivo)


@app.route('/exportar/curriculos')
@login_required
@rol_requerido('admin', 'empleado')
def exportar_curriculos_excel():
    lista = Postulante.query.order_by(Postulante.puntaje_ml.desc()).all()

    wb = Workbook()
    ws = wb.active
    ws.title = 'Postulantes'

    _encabezado_institucional(ws, 'REPORTE DE EVALUACIÓN DE CURRÍCULOS — ML')

    cols   = ['N°', 'Nombre', 'DNI', 'Puesto', 'Exp. (años)',
              'Educación', 'Habilidades', 'Idiomas', 'Resultado ML',
              'Puntaje ML', 'Probabilidad', 'Fecha']
    anchos = [4, 22, 10, 18, 10, 14, 11, 9, 12, 11, 12, 16]
    for i, w in enumerate(anchos, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    _fila_encabezado_tabla(ws, 6, cols)

    for idx, p in enumerate(lista):
        fila = 7 + idx
        datos = [
            idx + 1, p.nombre, p.dni or '—', p.puesto,
            p.experiencia, p.educacion.title(),
            p.habilidades, p.idiomas or '0',
            p.resultado_ml.upper(),
            f'{p.puntaje_ml}%',
            f'{round(p.probabilidad * 100, 1)}%',
            p.fecha_registro.strftime('%d/%m/%Y')
        ]
        _fila_datos(ws, fila, datos, idx % 2 == 1)
        color_r = 'FFE1F5EE' if p.resultado_ml == 'apto' else 'FFFCEBEB'
        ws.cell(fila, 9).fill = PatternFill('solid', fgColor=color_r)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    nombre_archivo = f'curriculos_{datetime.now().strftime("%Y%m%d_%H%M")}.xlsx'
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name=nombre_archivo)

@app.route('/auditoria')
@login_required
@rol_requerido('admin')
def auditoria():
    registros = Auditoria.query.order_by(Auditoria.fecha.desc()).all()
    return render_template('auditoria.html', registros=registros)

# ════════════════════════════════════════════════════════════════════
# INICIALIZAR BD
# ════════════════════════════════════════════════════════════════════
def inicializar_bd():
    with app.app_context():
        db.create_all()
        if not Usuario.query.first():
            usuarios_default = [
                Usuario(nombre='Administrador', username='admin',
                        password=generate_password_hash('admin123'),
                        rol='admin', email='admin@gestimuni.pe',
                        dni='00000001'),
                Usuario(nombre='María Quispe', username='empleado1',
                        password=generate_password_hash('emp123'),
                        rol='empleado', email='empleado@gestimuni.pe',
                        dni='12345678'),
                Usuario(nombre='Juan Pérez', username='ciudadano1',
                        password=generate_password_hash('1234'),
                        rol='ciudadano', email='ciudadano@gmail.com',
                        dni='87654321'),
                Usuario(nombre='Postulante Demo', username='postulante1',
                        password=generate_password_hash('post123'),
                        rol='postulante', email='postulante@gmail.com',
                        dni='11223344'),

                        Usuario(nombre='Aldahir', username='ADMIN',
                        password=generate_password_hash('ADMIN123'),
                        rol='admin', email='pasivi22@gmail.com',
                        dni='71870033'),
            ]
            db.session.add_all(usuarios_default)
            db.session.commit()
            print('[BD] Inicializada con usuarios por defecto.')  


# ════════════════════════════════════════════════════════════════════
# IA AVANZADA PARA EVALUACIÓN DE CURRÍCULOS + CONVOCATORIAS
# ════════════════════════════════════════════════════════════════════
def _normalizar_texto(texto):
    '''Minúsculas, sin tildes y espacios normalizados.'''
    if not texto:
        return ""
    texto = unicodedata.normalize("NFKD", str(texto))
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = re.sub(r"\s+", " ", texto)
    return texto.lower().strip()


def _ruta_data(nombre):
    carpeta = os.path.join(os.path.dirname(__file__), 'data')
    os.makedirs(carpeta, exist_ok=True)
    return os.path.join(carpeta, nombre)


def cargar_puestos():
    ruta = _ruta_data('puestos.json')
    if os.path.exists(ruta):
        try:
            with open(ruta, 'r', encoding='utf-8') as f:
                data = json.load(f)
                if isinstance(data, dict) and data:
                    return data
        except Exception as e:
            print('No se pudo cargar puestos.json:', e)

    # Respaldo interno por si el JSON no existe
    return {
        'abogado': {'nombre':'Abogado Municipal','educacion_clave':['derecho','abogado','juridico','legal','leyes'],'habilidades_clave':['contratos','normativa','expedientes','derecho administrativo','litigios','asesoria legal'],'certificaciones_clave':['colegiatura','cal','derecho administrativo','contrataciones del estado'],'min_experiencia':1},
        'project_manager': {'nombre':'Project Manager','educacion_clave':['management','administracion','maestria','master','product owner'],'habilidades_clave':['project manager','product manager','product owner','customer success','pmp','capm','trello','notion','analytics','zendesk','mixpanel','google analytics','tableau','intercom'],'certificaciones_clave':['pmp','capm','google ads'],'min_experiencia':2},
        'contador': {'nombre':'Contador Municipal','educacion_clave':['contabilidad','contador','finanzas','tributacion'],'habilidades_clave':['siga','siaf','presupuesto','balance','tesoreria','costos','excel'],'certificaciones_clave':['contabilidad gubernamental','siaf','siga'],'min_experiencia':1},
        'secretaria': {'nombre':'Secretaria Administrativa','educacion_clave':['secretariado','administracion','tecnico'],'habilidades_clave':['word','excel','archivo','redaccion','atencion al cliente','ofimatica'],'certificaciones_clave':['ofimatica','excel'],'min_experiencia':1}
    }


def _leer_json(nombre, defecto):
    ruta = _ruta_data(nombre)
    if not os.path.exists(ruta):
        return defecto
    try:
        with open(ruta, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception:
        return defecto


def _guardar_json(nombre, data):
    ruta = _ruta_data(nombre)
    with open(ruta, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def extraer_texto_cv(ruta, extension):
    '''Extrae texto real del CV. DOCX usa docx2txt + tablas + XML como respaldo.'''
    texto = ""
    try:
        if extension == 'pdf':
            reader = PdfReader(ruta)
            for page in reader.pages:
                texto += page.extract_text() or ""

        elif extension == 'docx':
            try:
                texto = docx2txt.process(ruta) or ""
            except Exception as e:
                print('docx2txt falló:', e)

            if not texto.strip():
                try:
                    doc = DocxDocument(ruta)
                    partes = []
                    partes.extend(p.text for p in doc.paragraphs if p.text.strip())
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    partes.append(cell.text.strip())
                    texto = "\n".join(partes)
                except Exception as e:
                    print('python-docx falló:', e)

            if not texto.strip():
                try:
                    with zipfile.ZipFile(ruta) as z:
                        xml = z.read('word/document.xml').decode('utf-8', errors='ignore')
                        texto = re.sub(r'<[^>]+>', ' ', xml)
                except Exception as e:
                    print('fallback XML DOCX falló:', e)

        elif extension in ['txt', 'csv']:
            with open(ruta, 'r', encoding='utf-8', errors='ignore') as f:
                texto = f.read()

    except Exception as e:
        print('ERROR EXTRAYENDO CV:', e)

    texto = re.sub(r'\s+', ' ', texto).strip()
    print('========== TEXTO CV EXTRAÍDO ==========', texto[:1200], '========== FIN TEXTO CV ==========', sep='\n')
    return texto


def _detectar_nombre(texto_original):
    secciones = {'FORMACION REGLADA','EXPERIENCIA PROFESIONAL','IDIOMAS','CERTIFICADOS','HARD SKILLS','QUIERES SABER MAS SOBRE MI'}
    candidatos = re.findall(r'\b([A-ZÁÉÍÓÚÑ]{3,}(?:\s+[A-ZÁÉÍÓÚÑ]{3,}){1,3})\b', texto_original)
    for c in candidatos:
        limpio = _normalizar_texto(c).upper()
        if limpio not in secciones and not any(s in limpio for s in ['FORMACION','EXPERIENCIA','CERTIFICADO','IDIOMA','HARD']):
            partes = c.title().split()
            # Evita duplicar nombre si el extractor lo repitió
            filtradas = []
            for p in partes:
                if not filtradas or filtradas[-1].lower() != p.lower():
                    filtradas.append(p)
            return ' '.join(filtradas[:3])
    return None


def _detectar_dni(texto_original):
    m = re.search(r'\b(?:dni|documento|identidad)\D{0,10}(\d{8})\b', texto_original, re.I)
    if m:
        return m.group(1)
    return None


def _detectar_fecha_nacimiento_y_edad(texto_original):
    patrones = [
        r'(?:fecha de nacimiento|nacimiento|fec\. nac\.?|f\. nac\.?)\D{0,20}(\d{1,2}[/-]\d{1,2}[/-]\d{4})',
        r'(?:fecha de nacimiento|nacimiento)\D{0,20}(\d{4}[/-]\d{1,2}[/-]\d{1,2})'
    ]
    for pat in patrones:
        m = re.search(pat, texto_original, re.I)
        if m:
            fecha = m.group(1)
            for fmt in ['%d/%m/%Y','%d-%m-%Y','%Y/%m/%d','%Y-%m-%d']:
                try:
                    dt = datetime.strptime(fecha, fmt)
                    hoy = datetime.now()
                    edad = hoy.year - dt.year - ((hoy.month, hoy.day) < (dt.month, dt.day))
                    return fecha, edad
                except Exception:
                    pass
    return None, None


def _experiencia_por_rangos(texto_lower):
    '''Calcula experiencia evitando duplicar rangos superpuestos.'''
    intervalos = []
    # Rangos tipo 02/2016 - Presente
    for mes_i, anio_i, fin in re.findall(r'(\d{1,2})/(\d{4})\s*-\s*(presente|actualidad|\d{1,2}/\d{4})', texto_lower):
        inicio = int(anio_i) * 12 + int(mes_i)
        if fin in ['presente', 'actualidad']:
            now = datetime.now()
            final = now.year * 12 + now.month
        else:
            mes_f, anio_f = fin.split('/')
            final = int(anio_f) * 12 + int(mes_f)
        if final > inicio:
            intervalos.append((inicio, final))

    if not intervalos:
        return 0

    intervalos = sorted(set(intervalos))
    fusionados = []
    for ini, fin in intervalos:
        if not fusionados or ini > fusionados[-1][1]:
            fusionados.append([ini, fin])
        else:
            fusionados[-1][1] = max(fusionados[-1][1], fin)
    meses = sum(fin - ini for ini, fin in fusionados)
    return round(meses / 12)


def analizar_cv_por_puesto(texto, puesto_key):
    puestos = cargar_puestos()
    perfil = puestos.get(puesto_key) or next(iter(puestos.values()))

    texto_original = texto or ''
    texto_lower = _normalizar_texto(texto_original)

    correo_match = re.search(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', texto_original)
    correo = correo_match.group(0) if correo_match else None

    telefono_match = re.search(r'(\+?\d[\d\s]{7,}\d)', texto_original)
    telefono = telefono_match.group(0).strip() if telefono_match else None

    nombre_detectado = _detectar_nombre(texto_original)
    dni_detectado = _detectar_dni(texto_original)
    fecha_nacimiento, edad = _detectar_fecha_nacimiento_y_edad(texto_original)

    # Experiencia: si el CV declara "15 años", se prioriza; si no, se calcula por rangos fusionados.
    declarados = [int(x[0]) for x in re.findall(r'(\d{1,2})\s*(anos|años|ano|año)\s+de\s+experiencia', texto_lower)]
    if not declarados:
        declarados = [int(x[0]) for x in re.findall(r'en\s+(\d{1,2})\s*(anos|años)', texto_lower)]
    exp_declarada = max(declarados) if declarados else 0
    exp_rangos = _experiencia_por_rangos(texto_lower)
    experiencia = exp_declarada or exp_rangos
    experiencia = max(0, min(experiencia, 40))  # límite razonable

    # Educación
    educacion = 'secundaria'
    educacion_detectada = []
    formaciones = {
        'Doctorado': ['doctorado', 'phd'],
        'Maestría / Master': ['maestria', 'master', 'mba'],
        'Universidad': ['universidad', 'bachiller', 'licenciado', 'titulo profesional'],
        'Técnico': ['tecnico', 'instituto']
    }
    for etiqueta, claves in formaciones.items():
        if any(c in texto_lower for c in claves):
            educacion_detectada.append(etiqueta)
    if any(x in texto_lower for x in ['doctorado', 'phd']):
        educacion = 'doctorado'
    elif any(x in texto_lower for x in ['maestria', 'master', 'mba', 'postgrado']):
        educacion = 'postgrado'
    elif any(x in texto_lower for x in ['universidad', 'universitario', 'bachiller', 'licenciado']):
        educacion = 'universitario'
    elif 'tecnico' in texto_lower:
        educacion = 'tecnico'

    # Idiomas
    mapa_idiomas = {'ingles':'Inglés','espanol':'Español','portugues':'Portugués','frances':'Francés','quechua':'Quechua'}
    idiomas_detectados = []
    for clave, etiqueta in mapa_idiomas.items():
        if clave in texto_lower and etiqueta not in idiomas_detectados:
            idiomas_detectados.append(etiqueta)

    # Certificaciones sin duplicados
    mapa_cert = {
        'pmp': 'PMP', 'pmps': 'PMP', 'capm': 'CAPM', 'capms': 'CAPM',
        'google ads': 'Google Ads', 'campanas de busqueda': 'Google Ads', 'campanas de shopping': 'Google Ads',
        'siga': 'SIGA', 'siaf': 'SIAF', 'seace': 'SEACE', 'scrum': 'Scrum', 'power bi': 'Power BI'
    }
    certificaciones_detectadas = []
    for clave, etiqueta in mapa_cert.items():
        if clave in texto_lower and etiqueta not in certificaciones_detectadas:
            certificaciones_detectadas.append(etiqueta)

    # Skills
    skills_generales = ['Tableau','Mixpanel','Google Analytics 360','Google Analytics','Notion','Trello','Zendesk','Intercom','Excel','Word','SQL','Python','Java','Power BI','AutoCAD','SIAF','SIGA','SEACE','PHP','JavaScript','MySQL','Flask','Laravel','React','Redacción','Archivo','Atención al cliente']
    hard_skills_detectadas = []
    for skill in skills_generales:
        if _normalizar_texto(skill) in texto_lower and skill not in hard_skills_detectadas:
            hard_skills_detectadas.append(skill)

    cargos_clave = ['Ecommerce Product Owner','Product Owner','Product Manager','Project Manager','Customer Success Manager','Customer Success','Agente de Servicio al cliente','Administrador','Abogado','Contador','Ingeniero Civil','Arquitecto','Programador','Secretaria','Asistente Administrativo']
    experiencia_relevante = []
    for cargo in cargos_clave:
        if _normalizar_texto(cargo) in texto_lower and cargo not in experiencia_relevante:
            experiencia_relevante.append(cargo)

    edu_match = sum(1 for p in perfil.get('educacion_clave', []) if _normalizar_texto(p) in texto_lower)
    hab_match = sum(1 for p in perfil.get('habilidades_clave', []) if _normalizar_texto(p) in texto_lower)
    cert_match = sum(1 for p in perfil.get('certificaciones_clave', []) if _normalizar_texto(p) in texto_lower)

    # Puntaje explicable 100
    score_exp = 25 if experiencia >= int(perfil.get('min_experiencia', 0)) else min(20, experiencia * 5)
    score_edu = min(20, edu_match * 10 + (10 if educacion in ['postgrado', 'doctorado'] else 0))
    score_hab = min(25, hab_match * 5 + min(len(hard_skills_detectadas), 10))
    score_cert = min(15, cert_match * 5 + len(certificaciones_detectadas) * 3)
    score_cargos = 10 if experiencia_relevante else 0
    score_idiomas = 5 if len(idiomas_detectados) >= 2 else (3 if idiomas_detectados else 0)

    puntaje = score_exp + score_edu + score_hab + score_cert + score_cargos + score_idiomas

    # Penalización fuerte si es un puesto legal y no hay evidencia jurídica.
    if puesto_key == 'abogado' and not any(x in texto_lower for x in ['derecho', 'abogado', 'juridico', 'legal', 'leyes', 'contratos']):
        puntaje = min(puntaje, 35)

    puntaje = int(min(max(puntaje, 0), 100))
    resultado = 'apto' if puntaje >= 70 else ('revisar' if puntaje >= 50 else 'no_apto')
    recomendacion = 'RECOMENDADO PARA ENTREVISTA' if puntaje >= 70 else ('REVISAR MANUALMENTE' if puntaje >= 50 else 'NO RECOMENDADO')

    habilidades = max(hab_match, len(hard_skills_detectadas))
    certificaciones = max(cert_match, len(certificaciones_detectadas))

    explicacion = [
        'ANÁLISIS INTELIGENTE DEL CURRÍCULUM VITAE',
        f"Nombre detectado: {nombre_detectado or 'No detectado'}.",
        f"DNI detectado: {dni_detectado or 'No detectado'}.",
        f"Correo detectado: {correo or 'No detectado'}.",
        f"Teléfono detectado: {telefono or 'No detectado'}.",
        f"Fecha de nacimiento detectada: {fecha_nacimiento or 'No detectada'}.",
        f"Edad estimada: {str(edad) if edad is not None else 'No detectada'}.",
        f"Experiencia estimada: {experiencia} años.",
        f"Nivel educativo detectado: {educacion}.",
    ]
    if educacion_detectada:
        explicacion.append('Formación detectada: ' + ', '.join(educacion_detectada) + '.')
    if idiomas_detectados:
        explicacion.append('Idiomas detectados: ' + ', '.join(idiomas_detectados) + '.')
    if certificaciones_detectadas:
        explicacion.append('Certificaciones detectadas: ' + ', '.join(certificaciones_detectadas) + '.')
    if hard_skills_detectadas:
        explicacion.append('Hard skills detectadas: ' + ', '.join(hard_skills_detectadas) + '.')
    if experiencia_relevante:
        explicacion.append('Experiencia relevante detectada: ' + ', '.join(experiencia_relevante) + '.')

    explicacion.extend([
        '',
        'DETALLE DEL PUNTAJE',
        f'Experiencia: {score_exp}/25.',
        f'Educación: {score_edu}/20.',
        f'Habilidades: {score_hab}/25.',
        f'Certificaciones: {score_cert}/15.',
        f'Cargos relacionados: {score_cargos}/10.',
        f'Idiomas: {score_idiomas}/5.',
        f"Compatibilidad con el puesto {perfil.get('nombre', puesto_key)}: {puntaje}%.",
        f'Recomendación IA: {recomendacion}.'
    ])

    if resultado == 'apto':
        explicacion.append(f"Conclusión: APTO para el puesto de {perfil.get('nombre', puesto_key)}.")
    elif resultado == 'revisar':
        explicacion.append(f"Conclusión: REVISAR MANUALMENTE. El perfil tiene coincidencias parciales con {perfil.get('nombre', puesto_key)}.")
    else:
        explicacion.append(f"Conclusión: NO APTO para el puesto de {perfil.get('nombre', puesto_key)}.")

    return {
        'nombre_detectado': nombre_detectado,
        'dni_detectado': dni_detectado,
        'correo_detectado': correo,
        'telefono_detectado': telefono,
        'fecha_nacimiento': fecha_nacimiento,
        'edad': edad,
        'experiencia': experiencia,
        'educacion': educacion,
        'habilidades': habilidades,
        'certificaciones': certificaciones,
        'puntaje': puntaje,
        'resultado': resultado,
        'probabilidad': puntaje / 100,
        'recomendacion': recomendacion,
        'explicacion': '\n'.join(explicacion)
    }


# ════════════════════════════════════════════════════════════════════
# CONVOCATORIAS: PLAZO, RANKING Y CORREOS AUTOMÁTICOS
# ════════════════════════════════════════════════════════════════════
def cargar_convocatorias():
    return _leer_json('convocatorias.json', [])


def guardar_convocatorias(data):
    _guardar_json('convocatorias.json', data)


def _parse_dt(valor):
    if not valor:
        return None
    for fmt in ['%Y-%m-%dT%H:%M', '%Y-%m-%d %H:%M:%S', '%Y-%m-%d %H:%M']:
        try:
            return datetime.strptime(valor, fmt)
        except Exception:
            pass
    return None


def _html_resultado_postulante(nombre, puesto_nombre, seleccionado, puntaje):
    if seleccionado:
        titulo = '¡Felicidades! Has sido seleccionado para la siguiente etapa'
        color = '#1D9E75'
        cuerpo = f'Tu CV obtuvo el mejor puntaje para el puesto de {puesto_nombre}. Puntaje IA: {puntaje}%.'
    else:
        titulo = 'Resultado de evaluación de currículo'
        color = '#A32D2D'
        cuerpo = f'Agradecemos tu postulación al puesto de {puesto_nombre}. En esta oportunidad no fuiste seleccionado. Puntaje IA: {puntaje}%.'
    return f'''
    <div style="font-family:Arial,sans-serif;max-width:560px;margin:auto;border:1px solid #ddd;border-radius:10px;overflow:hidden;">
      <div style="background:#2E7D4F;color:#fff;padding:18px;text-align:center;">
        <h2 style="margin:0;">GestiMuni Huánuco</h2>
        <p style="margin:4px 0 0;font-size:13px;">Evaluación inteligente de currículos</p>
      </div>
      <div style="padding:22px;">
        <h3 style="color:{color};margin-top:0;">{titulo}</h3>
        <p>Hola <strong>{nombre}</strong>,</p>
        <p>{cuerpo}</p>
        <p style="font-size:12px;color:#777;">Este correo fue generado automáticamente por el módulo de IA de GestiMuni.</p>
      </div>
    </div>'''


def enviar_correo_postulante(postulante, seleccionado, puesto_nombre):
    if not postulante.email:
        return False
    try:
        asunto = 'Resultado de evaluación de CV'
        html = _html_resultado_postulante(postulante.nombre, puesto_nombre, seleccionado, postulante.puntaje_ml)
        msg = Message(
            subject=f'GestiMuni — {asunto}',
            sender=f"GestiMuni <{app.config['MAIL_USERNAME']}>",
            recipients=[postulante.email],
            html=html
        )
        mail.send(msg)
        return True
    except Exception as e:
        print('Error enviando correo a postulante:', e)
        return False


def evaluar_convocatoria_y_enviar(conv):
    puestos = cargar_puestos()
    puesto_key = conv.get('puesto')
    puesto_nombre = puestos.get(puesto_key, {}).get('nombre', puesto_key)
    inicio = _parse_dt(conv.get('fecha_inicio'))
    cierre = _parse_dt(conv.get('fecha_cierre'))

    q = Postulante.query.filter_by(puesto=puesto_key)
    if inicio:
        q = q.filter(Postulante.fecha_registro >= inicio)
    if cierre:
        q = q.filter(Postulante.fecha_registro <= cierre)
    candidatos = q.order_by(Postulante.puntaje_ml.desc(), Postulante.fecha_registro.asc()).all()

    if not candidatos:
        conv['estado'] = 'cerrada_sin_postulantes'
        conv['finalizada'] = True
        conv['resultado'] = {'mensaje': 'No hubo postulantes dentro del plazo.'}
        return conv

    mejor = candidatos[0]
    enviados = []
    for p in candidatos:
        seleccionado = (p.id == mejor.id and p.resultado_ml in ['apto', 'revisar'])
        ok = enviar_correo_postulante(p, seleccionado, puesto_nombre)
        enviados.append({'postulante_id': p.id, 'email': p.email, 'seleccionado': seleccionado, 'correo_enviado': ok})

    conv['estado'] = 'finalizada'
    conv['finalizada'] = True
    conv['resultado'] = {
        'ganador_id': mejor.id,
        'ganador_nombre': mejor.nombre,
        'ganador_puntaje': mejor.puntaje_ml,
        'total_postulantes': len(candidatos),
        'correos': enviados,
        'fecha_evaluacion': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    registrar_auditoria(
        f"Finalizó convocatoria {puesto_nombre}. Ganador: {mejor.nombre} ({mejor.puntaje_ml}%).",
        'Convocatorias'
    )
    return conv


def procesar_convocatorias_vencidas():
    if not current_user.is_authenticated or current_user.rol != 'admin':
        return
    data = cargar_convocatorias()
    cambio = False
    ahora = datetime.now()
    for conv in data:
        if conv.get('finalizada'):
            continue
        cierre = _parse_dt(conv.get('fecha_cierre'))
        if cierre and ahora >= cierre:
            evaluar_convocatoria_y_enviar(conv)
            cambio = True
    if cambio:
        guardar_convocatorias(data)


@app.route('/convocatorias')
@login_required
@rol_requerido('admin')
def convocatorias():
    procesar_convocatorias_vencidas()
    return render_template('convocatorias.html', convocatorias=cargar_convocatorias(), puestos=cargar_puestos())


@app.route('/convocatorias/nueva', methods=['GET', 'POST'])
@login_required
@rol_requerido('admin')
def nueva_convocatoria():
    puestos = cargar_puestos()
    if request.method == 'POST':
        puesto = request.form.get('puesto')
        duracion_tipo = request.form.get('duracion_tipo', 'dias')
        duracion = int(request.form.get('duracion', 1))
        inicio = datetime.now()
        if duracion_tipo == 'minutos':
            from datetime import timedelta
            cierre = inicio + timedelta(minutes=duracion)
        elif duracion_tipo == 'horas':
            from datetime import timedelta
            cierre = inicio + timedelta(hours=duracion)
        else:
            from datetime import timedelta
            cierre = inicio + timedelta(days=duracion)

        data = cargar_convocatorias()
        conv = {
            'id': uuid.uuid4().hex[:10],
            'puesto': puesto,
            'fecha_inicio': inicio.strftime('%Y-%m-%d %H:%M:%S'),
            'fecha_cierre': cierre.strftime('%Y-%m-%d %H:%M:%S'),
            'estado': 'abierta',
            'finalizada': False,
            'creado_por': current_user.nombre,
            'resultado': {}
        }
        data.append(conv)
        guardar_convocatorias(data)
        registrar_auditoria(f"Creó convocatoria para {puestos.get(puesto, {}).get('nombre', puesto)} hasta {conv['fecha_cierre']}", 'Convocatorias')
        flash('Convocatoria creada correctamente.', 'success')
        return redirect(url_for('convocatorias'))
    return render_template('convocatoria_form.html', puestos=puestos)


@app.route('/convocatorias/finalizar/<conv_id>', methods=['POST'])
@login_required
@rol_requerido('admin')
def finalizar_convocatoria(conv_id):
    data = cargar_convocatorias()
    for conv in data:
        if conv.get('id') == conv_id:
            evaluar_convocatoria_y_enviar(conv)
            guardar_convocatorias(data)
            flash('Convocatoria finalizada. Se enviaron correos a los postulantes registrados.', 'success')
            return redirect(url_for('convocatorias'))
    flash('Convocatoria no encontrada.', 'danger')
    return redirect(url_for('convocatorias'))


@app.route('/convocatorias/eliminar/<conv_id>', methods=['POST'])
@login_required
@rol_requerido('admin')
def eliminar_convocatoria(conv_id):
    data = [c for c in cargar_convocatorias() if c.get('id') != conv_id]
    guardar_convocatorias(data)
    registrar_auditoria('Eliminó una convocatoria', 'Convocatorias')
    flash('Convocatoria eliminada.', 'success')
    return redirect(url_for('convocatorias'))


if __name__ == '__main__':
    inicializar_bd()
    app.run(debug=True)
